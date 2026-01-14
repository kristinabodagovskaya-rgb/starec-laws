from flask import Flask, render_template, render_template_string, request, send_file, redirect, jsonify
import psycopg2
import psycopg2.extras
import re
from bs4 import BeautifulSoup
import openai
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
from functools import lru_cache

app = Flask(__name__)

# Словарь сокращений законов (добавлено 13.01.2026)
LAW_ABBREVIATIONS = {
    'гк': 'гражданский кодекс',
    'гк рф': 'гражданский кодекс',
    'ук': 'уголовный кодекс',
    'ук рф': 'уголовный кодекс',
    'тк': 'трудовой кодекс',
    'тк рф': 'трудовой кодекс',
    'кап': 'кодекс административных правонарушений',
    'коап': 'кодекс административных правонарушений',
    'нк': 'налоговый кодекс',
    'нк рф': 'налоговый кодекс',
    'жк': 'жилищный кодекс',
    'жк рф': 'жилищный кодекс',
    'ск': 'семейный кодекс',
    'ск рф': 'семейный кодекс',
    'зк': 'земельный кодекс',
    'зк рф': 'земельный кодекс',
    'апк': 'арбитражный процессуальный кодекс',
    'гпк': 'гражданский процессуальный кодекс',
    'упк': 'уголовно-процессуальный кодекс',
    'бк': 'бюджетный кодекс',
    'вк': 'водный кодекс',
    'лк': 'лесной кодекс',
    'зпп': 'защите прав потребителей',
    'озпп': 'защите прав потребителей',
    'о защите прав': 'защите прав потребителей',
    'зозпп': 'защите прав потребителей',
}


def expand_abbreviations(query):
    """Раскрывает сокращения в поисковом запросе"""
    query_lower = query.lower().strip()
    for abbr, full in LAW_ABBREVIATIONS.items():
        if abbr in query_lower:
            query_lower = query_lower.replace(abbr, full)
    return query_lower

PG_CONFIG = {
    'host': '127.0.0.1',
    'port': 5432,
    'database': 'starec_laws',
    'user': 'flaskapp',
    'password': 'flaskpass123'
}

# OpenAI API key (should be in environment variable)
openai.api_key = os.getenv('OPENAI_API_KEY', '')

# Simple HTML template for law list (index page)
INDEX_TEMPLATE = '''
<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>База законов Российской Федерации</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body {
            font-family: 'Times New Roman', serif;
            line-height: 1.6;
            color: #000;
            background: #f5f5f5;
            padding: 20px;
        }
        .container {
            max-width: 900px;
            margin: 0 auto;
            background: white;
            padding: 40px 60px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }
        .header {
            text-align: center;
            margin-bottom: 40px;
            padding-bottom: 20px;
            border-bottom: 2px solid #000;
        }
        .header h1 {
            font-size: 24px;
            font-weight: bold;
            margin-bottom: 10px;
            text-transform: uppercase;
        }
        .law-meta {
            font-size: 14px;
            color: #666;
            margin-top: 10px;
        }
        .law-list {
            list-style: none;
        }
        .law-list li {
            padding: 10px;
            border-bottom: 1px solid #e0e0e0;
        }
        .law-list a {
            color: #1d1d1f;
            text-decoration: none;
            font-size: 16px;
        }
        .law-list a:hover {
            text-decoration: underline;
        }
        .search-box {
            margin: 30px 0;
            text-align: center;
        }
        .search-box input {
            width: 60%;
            padding: 12px 20px;
            font-size: 16px;
            border: 1px solid #ccc;
            border-radius: 4px;
        }
        .search-box button {
            padding: 12px 30px;
            font-size: 16px;
            background: #1d1d1f;
            color: white;
            border: none;
            border-radius: 4px;
            cursor: pointer;
            margin-left: 10px;
        }
        .search-box button:hover {
            background: #000000;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>База законов Российской Федерации</h1>
            <div class="law-meta">Всего законов с текстом: {{ total_laws }}</div>
        </div>

        <div class="search-box">
            <form action="/search" method="get">
                <input type="text" name="q" placeholder="Поиск" required>
                <button type="submit">Найти</button>
            </form>
        </div>

        <ul class="law-list">
        {% for l in laws %}
            <li><a href="/law/{{ l.id }}">{{ l.title }}</a></li>
        {% endfor %}
        </ul>
    </div>
</body>
</html>
'''


def fix_merged_numbers(html):
    """
    Исправляет слипшиеся номера пунктов в тексте.
    Пример: "2Граждане" -> "2. Граждане"
    Добавлено 13.01.2026
    """
    if not html:
        return html
    # Добавляет точку и пробел после номера пункта, если он слипся с русской буквой
    return re.sub(r'</a>(\d+)([А-ЯЁа-яё])', r'</a>\1. \2', html)


def parse_law_structure(full_text_html):
    """
    Parse law HTML structure to extract:
    - Sections (РАЗДЕЛ)
    - Chapters (Глава)
    - Articles (Статья)

    UPDATED 13.01.2026:
    - Добавлена поддержка формата Кодекса (class='H' для заголовков статей)
    - Используется soup.find_all вместо итерации по children

    Returns:
        toc: List of sections with chapters and articles
        articles: Flat list of all articles
        full_text_processed: HTML with article-section wrappers
    """
    if not full_text_html:
        return [], [], ""

    # Исправляем слипшиеся номера перед парсингом
    full_text_html = fix_merged_numbers(full_text_html)

    soup = BeautifulSoup(full_text_html, 'html.parser')

    toc = []
    articles = []
    current_section = None
    current_chapter = None
    article_counter = 0

    # Ищем статьи по class='H' (формат Кодекса) или по традиционным div с id
    article_headers = soup.find_all('p', class_='H')

    if article_headers:
        # Формат Кодекса: статьи помечены class='H'
        for header in article_headers:
            text = header.get_text(strip=True)

            # Проверяем тип заголовка
            if text.startswith('РАЗДЕЛ') or text.startswith('Раздел'):
                current_section = {
                    'title': text,
                    'chapters': [],
                    'articles': []
                }
                toc.append(current_section)
                current_chapter = None

            elif text.startswith('Глава') or text.startswith('ГЛАВА'):
                chapter = {
                    'title': text,
                    'articles': []
                }
                if current_section:
                    current_section['chapters'].append(chapter)
                    current_chapter = chapter
                else:
                    # Глава без раздела - создаём дефолтную секцию
                    current_section = {
                        'title': '',
                        'chapters': [chapter],
                        'articles': []
                    }
                    toc.append(current_section)
                    current_chapter = chapter

            elif text.startswith('Статья') or text.startswith('СТАТЬЯ'):
                article_counter += 1
                article_id = f'st{article_counter}'

                # Добавляем id к параграфу для навигации
                header['id'] = article_id
                header['class'] = header.get('class', []) + ['article-header']

                article_entry = {
                    'id': article_id,
                    'title': text[:100] + '...' if len(text) > 100 else text
                }

                articles.append(article_entry)

                if current_chapter:
                    current_chapter['articles'].append(article_entry)
                elif current_section:
                    current_section['articles'].append(article_entry)
                else:
                    # Статья без раздела/главы
                    if not toc:
                        toc.append({
                            'title': '',
                            'chapters': [],
                            'articles': []
                        })
                    toc[0]['articles'].append(article_entry)

        # Возвращаем обработанный HTML
        full_text_processed = str(soup)

    else:
        # Традиционный формат: div с id="stXXX"
        processed_html_parts = []
        root_elements = list(soup.children) if soup.name is None else [soup]

        def process_children(parent):
            nonlocal current_section, current_chapter

            for elem in parent.children:
                if isinstance(elem, str):
                    if elem.strip():
                        processed_html_parts.append(elem)
                    continue

                if not hasattr(elem, 'name'):
                    continue

                text = elem.get_text(strip=True)

                if elem.name == 'p' and text.startswith('РАЗДЕЛ'):
                    current_section = {
                        'title': text,
                        'chapters': [],
                        'articles': []
                    }
                    toc.append(current_section)
                    current_chapter = None
                    processed_html_parts.append(str(elem))

                elif elem.name == 'p' and text.startswith('Глава'):
                    if current_section:
                        current_chapter = {
                            'title': text,
                            'articles': []
                        }
                        current_section['chapters'].append(current_chapter)
                    processed_html_parts.append(str(elem))

                elif elem.name == 'div' and (elem.get('id', '').startswith('st') or elem.get('id', '').startswith('ст-') or 'article' in elem.get('class', [])):
                    article_id = elem.get('id')
                    article_title_elem = elem.find('h3')

                    if article_title_elem:
                        article_title = article_title_elem.get_text(strip=True)

                        articles.append({
                            'id': article_id,
                            'title': article_title
                        })

                        article_entry = {
                            'id': article_id,
                            'title': article_title
                        }

                        if current_chapter:
                            current_chapter['articles'].append(article_entry)
                        elif current_section:
                            current_section['articles'].append(article_entry)

                        elem['class'] = elem.get('class', []) + ['article-section']
                        processed_html_parts.append(str(elem))
                    else:
                        processed_html_parts.append(str(elem))

                else:
                    processed_html_parts.append(str(elem))

        for root in root_elements:
            if hasattr(root, 'children'):
                process_children(root)
            elif hasattr(root, 'name'):
                processed_html_parts.append(str(root))

        full_text_processed = ''.join(processed_html_parts)

    # Если статьи найдены, но нет структуры TOC
    if not toc and articles:
        toc = [{
            'title': 'Статьи',
            'chapters': [],
            'articles': articles
        }]

    return toc, articles, full_text_processed


def get_query_embedding(query_text):
    """Generate embedding for search query using OpenAI"""
    try:
        response = openai.embeddings.create(
            model="text-embedding-3-small",
            input=query_text
        )
        return response.data[0].embedding
    except Exception as e:
        print(f"Error generating embedding: {e}")
        return None


@app.route('/')
def index():
    """Clean homepage with search and link to laws database"""
    # Динамический счётчик законов (добавлено 13.01.2026)
    try:
        conn = psycopg2.connect(**PG_CONFIG)
        cursor = conn.cursor()
        cursor.execute("SELECT COUNT(*) FROM law_embeddings")
        law_count = cursor.fetchone()[0]
        cursor.close()
        conn.close()
    except Exception as e:
        print(f"Error getting law count: {e}")
        law_count = 110  # fallback
    return render_template('index_clean.html', law_count=law_count)


@app.route('/laws')
def laws_list():
    """Laws list with filters"""
    conn = psycopg2.connect(**PG_CONFIG)
    cursor = conn.cursor()

    # Get filter parameters
    law_type = request.args.get('law_type', '').strip()
    search = request.args.get('search', '').strip()

    # Build query
    query = """
        SELECT id, title, law_number, law_date, last_amendment_date
        FROM law_embeddings
        WHERE full_text IS NOT NULL AND length(full_text) > 500
    """
    params = []

    # Apply filters
    if law_type == 'Кодекс':
        query += " AND title ILIKE %s"
        params.append('%кодекс%')
    elif law_type == 'ФКЗ':
        query += " AND law_number LIKE %s"
        params.append('%-ФКЗ')
    elif law_type == 'ФЗ':
        query += " AND (law_number LIKE %s AND law_number NOT LIKE %s)"
        params.extend(['%-ФЗ', '%-ФКЗ'])

    if search:
        query += " AND title ILIKE %s"
        params.append(f'%{search}%')

    query += " ORDER BY title"

    cursor.execute(query, params)
    laws = []
    for row in cursor.fetchall():
        laws.append({
            'id': row[0],
            'title': row[1],
            'law_number': row[2],
            'law_date': row[3],
            'last_amendment_date': row[4]
        })

    cursor.close()
    conn.close()

    return render_template('laws_list.html', laws=laws, law_type=law_type, search=search)


@app.route('/search')
def search():
    """Semantic search using vector similarity"""
    query = request.args.get('q', '').strip()

    if not query:
        return render_template_string(INDEX_TEMPLATE, laws=[], total_laws=0)

    # Раскрываем сокращения в запросе (добавлено 13.01.2026)
    expanded_query = expand_abbreviations(query)

    # Generate query embedding
    query_embedding = get_query_embedding(expanded_query)

    if not query_embedding:
        return "Error generating search embedding", 500

    # Search using cosine similarity
    conn = psycopg2.connect(**PG_CONFIG)
    cursor = conn.cursor()

    cursor.execute("""
        SELECT
            id,
            title,
            1 - (embedding <=> %s::vector) AS similarity
        FROM law_embeddings
        WHERE embedding IS NOT NULL
        ORDER BY similarity DESC
        LIMIT 20
    """, (query_embedding,))

    results = cursor.fetchall()
    cursor.close()
    conn.close()

    # Format results
    laws = [{'id': row[0], 'title': f"{row[1]} (релевантность: {row[2]:.2f})"} for row in results]

    return render_template_string(INDEX_TEMPLATE, laws=laws, total_laws=len(laws))


@app.route('/law/<int:law_id>')
def show_law(law_id):
    """Show single law with Garant-style UI"""
    conn = psycopg2.connect(**PG_CONFIG)
    cursor = conn.cursor()
    cursor.execute("""
        SELECT id, title, authority, eo_number, full_text,
               law_number, law_date, last_amendment_date, last_amendment_info
        FROM law_embeddings
        WHERE id = %s
    """, (law_id,))
    row = cursor.fetchone()

    if not row:
        cursor.close()
        conn.close()
        return "Закон не найден", 404

    law = {
        'id': row[0],
        'title': row[1],
        'authority': row[2],
        'eo_number': row[3],
        'full_text': row[4],
        'law_number': row[5],
        'law_date': row[6],
        'last_amendment_date': row[7],
        'last_amendment_info': row[8]
    }

    # Get all editions from law_editions table
    editions = []
    try:
        cursor.execute("""
            SELECT id, revision_date, revision_description
            FROM law_editions
            WHERE law_id = %s
            ORDER BY revision_date DESC
        """, (law_id,))
        for ed_row in cursor.fetchall():
            editions.append({
                'id': ed_row[0],
                'date': ed_row[1],
                'description': ed_row[2] or ''
            })
    except Exception as e:
        print(f"Error fetching editions: {e}")

    cursor.close()
    conn.close()

    # Parse law structure
    toc, articles, full_text_processed = parse_law_structure(law['full_text'])
    law['full_text_processed'] = full_text_processed

    return render_template(
        'law_detail_garant.html',
        law=law,
        toc=toc,
        articles=articles,
        editions=editions,
        is_revision_view=False
    )


@app.route('/law/<int:law_id>/edition/<int:edition_id>')
def show_law_edition(law_id, edition_id):
    """Show specific edition of a law"""
    conn = psycopg2.connect(**PG_CONFIG)
    cursor = conn.cursor()

    # Get law basic info
    cursor.execute("""
        SELECT id, title, law_number, law_date
        FROM law_embeddings
        WHERE id = %s
    """, (law_id,))
    law_row = cursor.fetchone()

    if not law_row:
        cursor.close()
        conn.close()
        return "Закон не найден", 404

    # Get specific edition
    cursor.execute("""
        SELECT id, revision_date, revision_description, full_text
        FROM law_editions
        WHERE id = %s AND law_id = %s
    """, (edition_id, law_id))
    ed_row = cursor.fetchone()

    if not ed_row:
        cursor.close()
        conn.close()
        return "Редакция не найдена", 404

    law = {
        'id': law_row[0],
        'title': law_row[1],
        'law_number': law_row[2],
        'law_date': law_row[3],
        'full_text': ed_row[3],
        'last_amendment_date': ed_row[1],
        'last_amendment_info': ed_row[2]
    }

    # Get all editions for navigation
    editions = []
    cursor.execute("""
        SELECT id, revision_date, revision_description
        FROM law_editions
        WHERE law_id = %s
        ORDER BY revision_date DESC
    """, (law_id,))
    for r in cursor.fetchall():
        editions.append({
            'id': r[0],
            'date': r[1],
            'description': r[2] or ''
        })

    cursor.close()
    conn.close()

    # Parse law structure
    toc, articles, full_text_processed = parse_law_structure(law['full_text'])
    law['full_text_processed'] = full_text_processed

    return render_template(
        'law_detail_garant.html',
        law=law,
        toc=toc,
        articles=articles,
        editions=editions,
        current_edition_id=edition_id,
        is_revision_view=True
    )


@app.route('/law/<int:law_id>/revision/<int:revision_id>')
def show_law_revision(law_id, revision_id):
    """Redirect old revision URLs to new edition URLs"""
    return redirect(f'/law/{law_id}/edition/{revision_id}', code=301)


@app.route('/law/<int:law_id>/download')
def download_law(law_id):
    """Download law as DOCX file"""
    conn = psycopg2.connect(**PG_CONFIG)
    cursor = conn.cursor()
    cursor.execute("""
        SELECT id, title, full_text, law_number, law_date, last_amendment_date
        FROM law_embeddings
        WHERE id = %s
    """, (law_id,))
    row = cursor.fetchone()
    cursor.close()
    conn.close()

    if not row:
        return "Закон не найден", 404

    law = {
        'id': row[0],
        'title': row[1],
        'full_text': row[2],
        'law_number': row[3],
        'law_date': row[4],
        'last_amendment_date': row[5]
    }

    # Create DOCX document
    doc = Document()

    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run(law['title'])
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add metadata
    if law['law_number'] and law['law_date']:
        meta = doc.add_paragraph()
        meta_text = f"от {law['law_date']} N {law['law_number']}"
        if law['last_amendment_date']:
            meta_text += f" (ред. от {law['last_amendment_date']})"
        meta_run = meta.add_run(meta_text)
        meta_run.font.size = Pt(12)
        meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()  # Empty line

    # Parse HTML and add content
    if law['full_text']:
        soup = BeautifulSoup(law['full_text'], 'html.parser')

        for elem in soup.find_all(['p', 'h3', 'div']):
            text = elem.get_text(strip=True)
            if not text:
                continue

            # Check element type
            if elem.name == 'h3':
                # Article heading
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.bold = True
                run.font.size = Pt(14)
            elif text.startswith('РАЗДЕЛ'):
                # Section title
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.bold = True
                run.font.size = Pt(14)
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif text.startswith('Глава'):
                # Chapter title
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.bold = True
                run.font.size = Pt(13)
            else:
                # Regular paragraph
                para = doc.add_paragraph(text)
                para.style = 'Normal'

    # Save to BytesIO
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # Generate filename
    filename = f"{law['law_number']}_{law_id}.docx" if law['law_number'] else f"law_{law_id}.docx"

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/api/autocomplete')
def autocomplete():
    """
    Автоподбор для поиска законов.
    Добавлено 13.01.2026
    """
    query = request.args.get('q', '').strip()

    if not query or len(query) < 2:
        return jsonify([])

    # Раскрываем сокращения
    expanded_query = expand_abbreviations(query)

    try:
        conn = psycopg2.connect(**PG_CONFIG)
        cursor = conn.cursor()

        # Ищем законы по названию
        cursor.execute("""
            SELECT id, title
            FROM law_embeddings
            WHERE title ILIKE %s
            ORDER BY
                CASE WHEN title ILIKE %s THEN 0 ELSE 1 END,
                title
            LIMIT 10
        """, (f'%{expanded_query}%', f'{expanded_query}%'))

        results = []
        for row in cursor.fetchall():
            results.append({
                'id': row[0],
                'title': row[1],
                'url': f'/law/{row[0]}'
            })

        cursor.close()
        conn.close()

        return jsonify(results)

    except Exception as e:
        print(f"Autocomplete error: {e}")
        return jsonify([])


# ========== V2.0 ROUTES ==========

@app.route('/v2/')
def v2_index():
    conn = psycopg2.connect(**PG_CONFIG)
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    cur.execute("SELECT COUNT(*) as cnt FROM legal_regimes")
    regimes_count = cur.fetchone()['cnt']

    cur.execute("SELECT COUNT(*) as cnt FROM legal_states")
    states_count = cur.fetchone()['cnt']

    cur.execute("SELECT id, name, category, description FROM legal_regimes WHERE parent_regime_id IS NULL ORDER BY id")
    regimes = cur.fetchall()

    cur.close()
    conn.close()

    stats = {
        'regimes_count': regimes_count,
        'states_count': states_count,
        'situations_count': 0,
        'active_norms_count': 0
    }

    return render_template('v2/index.html', stats=stats, regimes=regimes)


@app.route('/v2/situations/')
def v2_situations():
    return render_template('v2/situations.html', situations=[])


@app.route('/v2/regimes/')
def v2_regimes():
    conn = psycopg2.connect(**PG_CONFIG)
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    cur.execute("SELECT id, name, category, description FROM legal_regimes ORDER BY name")
    regimes = cur.fetchall()
    cur.close()
    conn.close()
    return render_template('v2/regimes.html', regimes=regimes)


@app.route('/v2/timeline/')
def v2_timeline():
    return render_template('v2/timeline.html', changes=[])


@app.route('/v2/regime/<int:regime_id>')
def v2_regime_detail(regime_id):
    conn = psycopg2.connect(**PG_CONFIG)
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    cur.execute("SELECT id, name, category, description FROM legal_regimes WHERE id = %s", (regime_id,))
    regime = cur.fetchone()

    if not regime:
        cur.close()
        conn.close()
        return "Режим не найден", 404

    cur.execute("SELECT id, name, description FROM legal_states WHERE regime_id = %s ORDER BY name", (regime_id,))
    states = cur.fetchall()

    cur.close()
    conn.close()

    return render_template('v2/regime_detail.html', regime=regime, states=states)


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080)
